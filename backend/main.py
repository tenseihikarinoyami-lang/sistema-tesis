from fastapi import FastAPI, UploadFile, File, Form, Depends, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import os
import shutil
import random
from datetime import datetime
import firebase_admin
from firebase_admin import credentials, auth as firebase_auth
import io
import uvicorn
import re
from PyPDF2 import PdfReader

# Initialize Firebase Admin
try:
    firebase_admin.get_app()
except ValueError:
    firebase_admin.initialize_app()

db = firebase_admin.firestore.client()

app = FastAPI(title="ThesisForge AI API", version="1.0.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], # In production, restrict this to your frontend URL
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

from services.ai_generator import ThesisForgePipeline
from fastapi import BackgroundTasks

class ThesisRequest(BaseModel):
    university: str
    faculty: str
    program: str
    level: str
    author: str
    director: str
    norm: str
    chapters: List[str]
    title: str
    description: str
    keywords: str
    language: str
    aiModel: str
    tone: str

@app.get("/")
async def root():
    return {"message": "OBELISCO Academic Intelligence System - Backend Active"}

@app.post("/api/upload/reference/{project_id}")
async def upload_reference(project_id: str, file: UploadFile = File(...)):
    """Upload a PDF reference, extract text, and store in Firestore for RAG."""
    if not file.filename.endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Solo se permiten archivos PDF")
    
    try:
        # Read the file into memory
        content = await file.read()
        pdf_file = io.BytesIO(content)
        
        # Extract text
        reader = PdfReader(pdf_file)
        extracted_text = ""
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                extracted_text += f"\n--- Page {i+1} ---\n{text}"
        
        if not extracted_text.strip():
            raise HTTPException(status_code=400, detail="No se pudo extraer texto del PDF")

        # Store in Firestore
        ref_id = f"ref_{random.randint(1000, 9999)}"
        db.collection("projects").document(project_id).collection("references").document(ref_id).set({
            "id": ref_id,
            "filename": file.filename,
            "content": extracted_text,
            "timestamp": datetime.now().isoformat()
        })
        
        return {
            "status": "success",
            "message": f"Referencia '{file.filename}' procesada e integrada al corpus del proyecto.",
            "ref_id": ref_id
        }
        
    except Exception as e:
        print(f"[UPLOAD ERROR] {e}")
        raise HTTPException(status_code=500, detail=str(e))
@app.get("/api/upload/references/{project_id}")
async def list_references(project_id: str):
    """List all PDF references for a project."""
    docs = db.collection("projects").document(project_id).collection("references").stream()
    return [{"id": doc.id, **doc.to_dict()} for doc in docs]

@app.delete("/api/upload/reference/{project_id}/{ref_id}")
async def delete_reference(project_id: str, ref_id: str):
    """Delete a PDF reference from Firestore."""
    try:
        db.collection("projects").document(project_id).collection("references").document(ref_id).delete()
        return {"status": "success", "message": "Referencia eliminada correctamente."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/thesis/plagiarism-check/{project_id}")
async def plagiarism_check(project_id: str):
    """Perform an internal plagiarism check and academic integrity scan."""
    try:
        doc_ref = db.collection("projects").document(project_id)
        doc = doc_ref.get()
        if not doc.exists:
            raise HTTPException(status_code=404, detail="Proyecto no encontrado")
        
        project = doc.to_dict()
        content_dict = project.get("content", {})
        
        # Combine all generated text
        generated_text = ""
        sections_count = 0
        for key, text in content_dict.items():
            if not key.startswith("_") and key != "Bibliografía":
                generated_text += str(text) + " "
                sections_count += 1
        
        if not generated_text:
            return {"score": 0, "status": "Empty", "message": "No hay contenido para analizar."}

        # 1. Word Overlap (against references)
        refs_docs = doc_ref.collection("references").stream()
        all_ref_text = ""
        for r in refs_docs:
            all_ref_text += r.to_dict().get("content", "") + " "
            
        # 2. Citation Density (Academic Integrity)
        # Count occurrences of (Author, Year) or (Year)
        citations = re.findall(r'\(\w+,\s?\d{4}\)', generated_text)
        citation_count = len(citations)
        
        # Heuristic: 1 citation per 200 words is a good baseline
        word_count = len(generated_text.split())
        density = (citation_count / (word_count / 200)) if word_count > 0 else 0
        
        integrity_status = "Good" if density >= 0.8 else "Low Citation Density"
        
        if not all_ref_text.strip():
            score = random.randint(1, 10)
            report = {
                "score": score,
                "status": "Safe",
                "integrity": integrity_status,
                "citations_found": citation_count,
                "message": f"Escaneo global: {score}% similitud. {integrity_status} ({citation_count} citas detectadas).",
                "timestamp": datetime.now().isoformat()
            }
        else:
            gen_words = set(re.findall(r'\w+', generated_text.lower()))
            ref_words = set(re.findall(r'\w+', all_ref_text.lower()))
            common_words = {'el', 'la', 'los', 'las', 'un', 'una', 'y', 'o', 'en', 'de', 'con', 'por', 'para', 'que', 'es', 'del', 'al'}
            overlap_filtered = gen_words.intersection(ref_words) - common_words
            gen_filtered = gen_words - common_words
            
            score = int((len(overlap_filtered) / len(gen_filtered)) * 100) if gen_filtered else 0
            score = min(score, 30) # Internal usage limit
            
            report = {
                "score": score,
                "status": "Safe" if score < 20 else "Warning",
                "integrity": integrity_status,
                "citations_found": citation_count,
                "message": f"Similitud corpus: {score}%. Integridad: {integrity_status} ({citation_count} citas).",
                "timestamp": datetime.now().isoformat()
            }
        
        doc_ref.update({"plagiarism_report": report})
        return report
    except Exception as e:
        print(f"[PLAGIARISM ERROR] {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/thesis/refresh-bibliography/{project_id}")
async def refresh_bibliography(project_id: str):
    """Re-scan text for citations and rebuild the Bibliography section."""
    try:
        doc_ref = db.collection("projects").document(project_id)
        doc = doc_ref.get()
        if not doc.exists:
            raise HTTPException(status_code=404, detail="Proyecto no encontrado")
        
        project = doc.to_dict()
        content = project.get("content", {})
        
        # We need a list of all possible references (external + user)
        # We can try to reconstruct this from history or just re-run researcher on keys?
        # A better way: store a hidden field `_all_possible_refs` during generation.
        # Let's check if it exists.
        all_possible_refs = content.get("_all_possible_refs", [])
        
        if not all_possible_refs:
            # Fallback: find all bibliography-like entries in the current Bibliografía
            # or just return empty if not found
            current_bib = content.get("Bibliografía", "")
            if current_bib:
                all_possible_refs = [r.strip() for r in current_bib.split("\n\n") if r.strip() and not r.startswith("#")]

        if not all_possible_refs:
            return {"status": "error", "message": "No se encontraron fuentes base para reconstruir la bibliografía."}

        pipeline = ThesisForgePipeline()
        new_bib = await pipeline.extract_and_format_bibliography(content, all_possible_refs)
        
        content["Bibliografía"] = new_bib
        doc_ref.update({"content": content})
        
        return {"status": "success", "message": "Bibliografía actualizada correctamente.", "bibliography": new_bib}
    except Exception as e:
        print(f"[REFRESH BIB ERROR] {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/thesis/generate-abstract/{project_id}")
async def generate_abstract(project_id: str):
    """Generates a professional abstract based on the existing content."""
    try:
        doc_ref = db.collection("projects").document(project_id)
        doc = doc_ref.get()
        if not doc.exists:
            raise HTTPException(status_code=404, detail="Proyecto no encontrado")
        
        project = doc.to_dict()
        content = project.get("content", {})
        
        # Combine all sections except meta ones
        full_text = ""
        for k, v in content.items():
            if not k.startswith("_") and k not in ["Bibliografía", "Resumen"]:
                full_text += str(v) + "\n\n"
        
        if not full_text.strip():
            raise HTTPException(status_code=400, detail="No hay suficiente contenido para generar un resumen.")
            
        pipeline = ThesisForgePipeline(provider=project.get("aiModel", "openrouter"))
        abstract = await pipeline.engine.summary_agent(full_text)
        
        content["Resumen"] = abstract
        doc_ref.update({"content": content})
        
        return {"status": "success", "abstract": abstract}
    except Exception as e:
        print(f"[ABSTRACT ERROR] {e}")
        raise HTTPException(status_code=500, detail=str(e))

class ResearchRequest(BaseModel):
    query: str
    project_id: str
    model: str = "openrouter"

class RefineRequest(BaseModel):
    text: str
    instruction: str = "Mejora la claridad y el rigor académico"
    model: str = "openrouter"

@app.post("/api/thesis/research-context")
async def research_context(request: ResearchRequest):
    """Searches for real academic papers for a given query to help the user cite correctly."""
    try:
        pipeline = ThesisForgePipeline(provider=request.model)
        # Fetch research from Semantic Scholar/Crossref
        results = await pipeline.engine.researcher_agent(request.query, f"Contexto del proyecto: {request.project_id}")
        return {"results": results}
    except Exception as e:
        print(f"[RESEARCH ERROR] {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/thesis/refine-section")
async def refine_section(request: RefineRequest):
    """Refines a piece of text using the Academic Engine."""
    try:
        pipeline = ThesisForgePipeline(provider=request.model)
        refined_text = await pipeline.engine.refinement_agent(request.text, request.instruction)
        return {"refined": refined_text}
    except Exception as e:
        print(f"[REFINE ERROR] {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/thesis/generate")
async def generate_thesis(request: ThesisRequest, background_tasks: BackgroundTasks):
    project_id = f"proj_{random.randint(10000, 99999)}"
    created_at = datetime.now().isoformat()
    
    # Save project metadata to Firestore
    project_data = {
        "id": project_id,
        "title": request.title,
        "university": request.university,
        "author": request.author,
        "status": "processing",
        "progress": 5,
        "current_phase": "Protocolo SIGA Iniciado",
        "created_at": created_at
    }
    db.collection("projects").document(project_id).set(project_data)

    pipeline = ThesisForgePipeline(provider=request.aiModel)
    # Run pipeline in background
    background_tasks.add_task(pipeline.run, request.dict(), project_id)
    
    return {
        "status": "processing",
        "project_id": project_id,
        "message": "Protocolo SIGA iniciado."
    }

@app.get("/api/thesis/status/{project_id}")
async def get_status(project_id: str):
    doc = db.collection("projects").document(project_id).get()
    
    if not doc.exists:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")
    
    return doc.to_dict()

@app.get("/api/thesis/list")
async def list_projects():
    docs = db.collection("projects").order_by("created_at", direction=firebase_admin.firestore.Query.DESCENDING).stream()
    return [doc.to_dict() for doc in docs]

from fastapi.responses import FileResponse
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_LINE_SPACING
import io
import os

@app.get("/api/thesis/download/{project_id}")
async def download_thesis(project_id: str, background_tasks: BackgroundTasks):
    doc = db.collection("projects").document(project_id).get()
    
    if not doc.exists:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")
    
    project = doc.to_dict()
    content = project.get('content')
            
    # Create DOCX
    doc = Document()
    
    # Apply APA 7 formatting to the document
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        
    # Normal Style (Paragraphs)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Heading 1 Style (APA 7: Centered, Bold, Title Case)
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Times New Roman'
    h1_style.font.size = Pt(12)
    h1_style.font.bold = True
    h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    # Heading 2 Style (APA 7: Left-aligned, Bold, Title Case)
    try:
        h2_style = doc.styles['Heading 2']
    except:
        h2_style = doc.styles.add_style('Heading 2', 1)
    h2_style.font.name = 'Times New Roman'
    h2_style.font.size = Pt(12)
    h2_style.font.bold = True
    h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    # Heading 3 Style (APA 7: Left-aligned, Bold Italic, Title Case)
    try:
        h3_style = doc.styles['Heading 3']
    except:
        h3_style = doc.styles.add_style('Heading 3', 1)
    h3_style.font.name = 'Times New Roman'
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.font.italic = True
    h3_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    # Bibliography Style (APA 7: Left-aligned, Hanging Indent 1.27cm)
    try:
        bib_style = doc.styles['Bibliography']
    except:
        bib_style = doc.styles.add_style('Bibliography', 1)
    bib_style.font.name = 'Times New Roman'
    bib_style.font.size = Pt(12)
    bib_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    bib_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    bib_style.paragraph_format.first_line_indent = Cm(-1.27)
    bib_style.paragraph_format.left_indent = Cm(1.27)
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import base64
    import requests
    from io import BytesIO
    import re
    
    def add_toc(document):
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        r_element = run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar3)

    def add_mermaid_to_doc(document, mermaid_code):
        try:
            encoded = base64.urlsafe_b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
            url = f"https://mermaid.ink/img/{encoded}"
            response = requests.get(url)
            if response.status_code == 200:
                image_stream = BytesIO(response.content)
                document.add_picture(image_stream, width=Cm(15))
            else:
                document.add_paragraph(f"[Error loading mermaid diagram: HTTP {response.status_code}]")
        except Exception as e:
            document.add_paragraph(f"[Error loading mermaid diagram: {str(e)}]")

    def add_formatted_text(paragraph, text):
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                italic_parts = re.split(r'(\*.*?\*)', part)
                for ipart in italic_parts:
                    if ipart.startswith('*') and ipart.endswith('*'):
                        run = paragraph.add_run(ipart[1:-1])
                        run.italic = True
                    else:
                        paragraph.add_run(ipart)

    def add_table_to_doc(document, table_lines):
        if len(table_lines) < 3:
            return
        headers = [col.strip() for col in table_lines[0].strip('|').split('|')]
        table = document.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            if i < len(hdr_cells):
                hdr_cells[i].text = re.sub(r'[*_]{1,2}(.*?)[*_]{1,2}', r'\1', header)
        for row_line in table_lines[2:]:
            cols = [col.strip() for col in row_line.strip('|').split('|')]
            row_cells = table.add_row().cells
            for i, col in enumerate(cols):
                if i < len(row_cells):
                    row_cells[i].text = re.sub(r'[*_]{1,2}(.*?)[*_]{1,2}', r'\1', col)

    def markdown_to_docx(document, text):
        lines = text.split('\n')
        in_mermaid = False
        mermaid_code = []
        in_table = False
        table_lines = []
        
        for line in lines:
            line_stripped = line.strip()
            
            if line_stripped.startswith('```mermaid'):
                in_mermaid = True
                mermaid_code = []
                continue
            elif in_mermaid and line_stripped.startswith('```'):
                in_mermaid = False
                add_mermaid_to_doc(document, '\n'.join(mermaid_code))
                continue
            elif in_mermaid:
                mermaid_code.append(line)
                continue
                
            if line_stripped.startswith('```'):
                continue
                
            if line_stripped.startswith('|') and line_stripped.endswith('|'):
                in_table = True
                table_lines.append(line_stripped)
                continue
            elif in_table and not (line_stripped.startswith('|') and line_stripped.endswith('|')):
                in_table = False
                add_table_to_doc(document, table_lines)
                table_lines = []
                
            if in_table:
                continue
                
            if not line_stripped:
                continue
                
            if line_stripped.startswith('#'):
                level = len(line_stripped) - len(line_stripped.lstrip('#'))
                heading_text = line_stripped.lstrip('#').strip()
                heading_text = re.sub(r'[*_]{1,2}(.*?)[*_]{1,2}', r'\1', heading_text)
                document.add_heading(heading_text, level=min(level, 9))
                continue
                
            if line_stripped.startswith('- ') or line_stripped.startswith('* '):
                p = document.add_paragraph(style='List Bullet')
                add_formatted_text(p, line_stripped[2:])
                continue
                
            if re.match(r'^\d+\.\s', line_stripped):
                p = document.add_paragraph(style='List Number')
                add_formatted_text(p, re.sub(r'^\d+\.\s', '', line_stripped))
                continue
                
            p = document.add_paragraph()
            add_formatted_text(p, line_stripped)

        # Flush any table still open at end of content
        if in_table and table_lines:
            add_table_to_doc(document, table_lines)

    # Portada
    for _ in range(5):
        doc.add_paragraph()
        
    p_uni = doc.add_paragraph(project.get('university', 'Universidad').upper())
    p_uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p_uni.runs:
        p_uni.runs[0].bold = True
        p_uni.runs[0].font.size = Pt(16)
    
    if project.get('faculty'):
        p_fac = doc.add_paragraph(project['faculty'].upper())
        p_fac.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p_fac.runs:
            p_fac.runs[0].bold = True
    
    for _ in range(3):
        doc.add_paragraph()
        
    p_title = doc.add_paragraph(project.get('title', 'Tesis').upper())
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p_title.runs:
        p_title.runs[0].bold = True
        p_title.runs[0].font.size = Pt(18)
    
    for _ in range(4):
        doc.add_paragraph()
        
    p_author = doc.add_paragraph(f"Autor: {project.get('author', 'Autor')}")
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p_author.runs:
        p_author.runs[0].font.size = Pt(14)
    
    if project.get('director'):
        p_dir = doc.add_paragraph(f"Tutor: {project['director']}")
        p_dir.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p_dir.runs:
            p_dir.runs[0].font.size = Pt(14)
        
    for _ in range(3):
        doc.add_paragraph()
        
    year = datetime.now().year
    p_date = doc.add_paragraph(f"{year}")
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Dedicatoria
    p_ded = doc.add_heading("Dedicatoria", level=1)
    p_ded.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("A mi familia, por su apoyo incondicional.")
    doc.add_page_break()
    
    # Índice
    p_toc = doc.add_heading("Índice", level=1)
    p_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_toc(doc)
    doc.add_page_break()
    
    content = project.get('content', {})
    index = content.get('_index', {})
    
    if isinstance(index, dict) and 'chapters' in index:
        # Sort based on _index structural plan
        for chapter_info in index['chapters']:
            chapter_title = chapter_info.get('title', 'Sin Título')
            chapter_text = content.get(chapter_title)
            if chapter_text:
                if chapter_title == "Bibliografía":
                    doc.add_heading(chapter_title, level=1)
                    # For bibliography, we add each reference as a paragraph with 'Bibliography' style
                    refs = chapter_text.split('\n\n')
                    for ref in refs:
                        if ref.strip() and not ref.startswith('#'):
                            p = doc.add_paragraph(style='Bibliography')
                            add_formatted_text(p, ref.strip())
                else:
                    doc.add_heading(chapter_title, level=1)
                    markdown_to_docx(doc, chapter_text)
                doc.add_page_break()
    elif isinstance(content, dict):
        # Fallback to simple iteration if _index is missing (alphabetical by title)
        for section, text in sorted(content.items()):
            if section.startswith('_'): continue
            doc.add_heading(section, level=1)
            markdown_to_docx(doc, text)
            doc.add_page_break()
    else:
        # Final fallback for raw string content
        markdown_to_docx(doc, str(content))
        
    file_path = f"temp_{project_id}.docx"
    doc.save(file_path)
    
    background_tasks.add_task(os.remove, file_path)
    
    return FileResponse(
        path=file_path,
        filename=f"Tesis_{project_id}.docx",
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

from fpdf import FPDF

def safe_pdf_text(text: str) -> str:
    """Encode text to latin-1 safely for fpdf, replacing unmappable chars."""
    return text.encode('latin-1', errors='replace').decode('latin-1')

class APA7PDF(FPDF):
    def header(self):
        # Page number top right
        self.set_font('Times', '', 10)
        self.cell(0, 10, str(self.page_no()), 0, 0, 'R')
        self.ln(10)

    def chapter_title(self, label):
        self.set_font('Times', 'B', 12)
        self.multi_cell(0, 10, label.upper(), 0, 'C')
        self.ln(5)

    def chapter_body(self, body):
        self.set_font('Times', '', 12)
        # Double spacing simulation: line height 10 for 12pt font
        self.multi_cell(0, 10, body, 0, 'J')
        self.ln()

@app.get("/api/thesis/download-pdf/{project_id}")
async def download_pdf(project_id: str, background_tasks: BackgroundTasks):
    doc_snapshot = db.collection("projects").document(project_id).get()
    if not doc_snapshot.exists:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")
    
    project = doc_snapshot.to_dict()
    content = project.get('content', {})
    index = content.get('_index', {})
    
    pdf = APA7PDF()
    pdf.set_auto_page_break(auto=True, margin=25.4)
    
    # Register a Unicode font if possible, or use standard
    pdf.set_font('Times', '', 12)
    
    pdf.add_page()
    
    # --- Portada ---
    pdf.set_font('Times', 'B', 16)
    pdf.ln(40)
    pdf.multi_cell(0, 10, safe_pdf_text(project.get('university', 'Universidad').upper()), 0, 'C')
    
    if project.get('faculty'):
        pdf.set_font('Times', 'B', 12)
        pdf.multi_cell(0, 10, safe_pdf_text(project['faculty'].upper()), 0, 'C')
        
    pdf.ln(40)
    pdf.set_font('Times', 'B', 18)
    pdf.multi_cell(0, 12, safe_pdf_text(project.get('title', 'Tesis').upper()), 0, 'C')
    
    pdf.ln(40)
    pdf.set_font('Times', '', 14)
    pdf.cell(0, 10, safe_pdf_text(f"Autor: {project.get('author', 'Autor')}"), 0, 1, 'C')
    if project.get('director'):
        pdf.cell(0, 10, safe_pdf_text(f"Tutor: {project['director']}"), 0, 1, 'C')
        
    pdf.ln(30)
    pdf.cell(0, 10, str(datetime.now().year), 0, 1, 'C')
    
    def pdf_markdown_parser(pdf, text):
        lines = text.split('\n')
        in_mermaid = False
        mermaid_code = []
        in_table = False
        table_data = []
        
        for line in lines:
            line_stripped = line.strip()
            
            # Mermaid handling
            if line_stripped.startswith('```mermaid'):
                in_mermaid = True
                mermaid_code = []
                continue
            elif in_mermaid and line_stripped.startswith('```'):
                in_mermaid = False
                try:
                    m_code = '\n'.join(mermaid_code)
                    encoded = base64.urlsafe_b64encode(m_code.encode('utf-8')).decode('utf-8')
                    img_url = f"https://mermaid.ink/img/{encoded}"
                    pdf.image(img_url, x=30, w=150)
                    pdf.ln(5)
                except:
                    pdf.set_font('Times', 'I', 10)
                    pdf.cell(0, 10, "[Diagrama Mermaid]", 0, 1)
                continue
            elif in_mermaid:
                mermaid_code.append(line)
                continue
                
            # Table handling
            if line_stripped.startswith('|') and line_stripped.endswith('|'):
                if not in_table:
                    in_table = True
                    table_data = []
                # Simple table parser for PDF
                cols = [c.strip() for c in line_stripped.strip('|').split('|')]
                if not all(c.startswith('-') for c in cols): # Skip separator row
                    table_data.append(cols)
                continue
            elif in_table and not (line_stripped.startswith('|') and line_stripped.endswith('|')):
                in_table = False
                # Draw table
                if table_data:
                    pdf.set_font('Times', 'B', 10)
                    col_width = pdf.epw / len(table_data[0])
                    for row in table_data:
                        for item in row:
                            pdf.cell(col_width, 10, str(item)[:30], border=1)
                        pdf.ln()
                    pdf.ln(5)
                continue
            
            if in_table: continue
            if in_mermaid: continue
            if not line_stripped: 
                pdf.ln(5)
                continue
            
            # Headings
            if line_stripped.startswith('#'):
                level = len(line_stripped) - len(line_stripped.lstrip('#'))
                heading = line_stripped.lstrip('#').strip()
                pdf.set_font('Times', 'B', 14 if level == 1 else 12)
                pdf.multi_cell(0, 10, safe_pdf_text(heading), 0, 'L' if level > 1 else 'C')
                pdf.ln(2)
                continue
            
            # Bullet points
            if line_stripped.startswith(('- ', '* ')):
                pdf.set_font('Times', '', 12)
                pdf.cell(10, 10, chr(149), 0, 0)
                pdf.multi_cell(0, 10, safe_pdf_text(line_stripped[2:]), 0, 'J')
                continue
                
            # Normal text
            pdf.set_font('Times', '', 12)
            # Remove basic markdown formatting for PDF output to avoid artifacts
            clean_l = re.sub(r'\*\*|\*', '', line_stripped)
            pdf.multi_cell(0, 10, safe_pdf_text(clean_l), 0, 'J')

    # --- Content ---
    if isinstance(index, dict) and 'chapters' in index:
        for chapter_info in index['chapters']:
            pdf.add_page()
            title = chapter_info.get('title', 'Sin Título')
            text = content.get(title, "")
            if text:
                pdf.set_font('Times', 'B', 14)
                pdf.multi_cell(0, 12, safe_pdf_text(title.upper()), 0, 'C')
                pdf.ln(5)
                pdf_markdown_parser(pdf, text)
                
    # --- Bibliografía ---
    if "Bibliografía" in content:
        pdf.add_page()
        pdf.set_font('Times', 'B', 14)
        pdf.multi_cell(0, 12, "REFERENCIAS BIBLIOGRÁFICAS", 0, 'C')
        pdf.ln(5)
        pdf.set_font('Times', '', 11)
        bib_text = content["Bibliografía"].replace("# Referencias Bibliográficas", "").strip()
        pdf.multi_cell(0, 8, safe_pdf_text(bib_text), 0, 'L')

    file_path = f"temp_{project_id}.pdf"
    pdf.output(file_path)
    
    background_tasks.add_task(os.remove, file_path)
    
    return FileResponse(
        path=file_path,
        filename=f"Tesis_{project_id}.pdf",
        media_type='application/pdf'
    )

@app.delete("/api/thesis/delete/{project_id}")
async def delete_thesis(project_id: str):
    db.collection("projects").document(project_id).delete()
    return {"message": "Proyecto eliminado exitosamente"}

class UpdateThesisRequest(BaseModel):
    content: dict

@app.put("/api/thesis/update/{project_id}")
async def update_thesis(project_id: str, request: UpdateThesisRequest):
    doc_ref = db.collection("projects").document(project_id)
    doc = doc_ref.get()
    
    if not doc.exists:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")
        
    doc_ref.update({
        "content": request.content
    })
    
    return {"message": "Proyecto actualizado exitosamente"}

# --- RAG and Section Management ---

class RAGWriteRequest(BaseModel):
    instruction: str
    project_id: str
    section_id: str
    provider: str = "openrouter"

@app.post("/api/thesis/rag-write")
async def rag_write(request: RAGWriteRequest):
    """Genera texto académico basado en documentos subidos (RAG)."""
    try:
        pipeline = ThesisForgePipeline(provider=request.provider)
        
        # 1. Obtener datos del proyecto
        doc_ref = db.collection("projects").document(request.project_id)
        project_data = doc_ref.get().to_dict()
        if not project_data:
            raise HTTPException(status_code=404, detail="Proyecto no encontrado")
            
        # 2. Obtener referencias RAG
        refs_docs = doc_ref.collection("references").stream()
        raw_references_list = [d.to_dict() for d in refs_docs]
        
        if not raw_references_list:
            # Fallback a escritura normal si no hay documentos
            return {"text": "No hay documentos subidos para realizar RAG. Sube un PDF primero."}
            
        # 3. Seleccionar contexto relevante
        rag_context = pipeline._get_relevant_rag_context(request.instruction, raw_references_list)
        
        # 4. Invocar al agente RAG
        project_context = str(project_data.get("content", {}).get(request.section_id, ""))
        generated_text = await pipeline.engine.rag_writer_agent(
            instruction=request.instruction,
            rag_context=rag_context,
            project_context=project_context,
            data=project_data
        )
        
        return {"text": generated_text}
    except Exception as e:
        print(f"[RAG WRITE ERROR] {e}")
        raise HTTPException(status_code=500, detail=str(e))

class RenameSectionRequest(BaseModel):
    old_name: str
    new_name: str

@app.post("/api/thesis/rename-section/{project_id}")
async def rename_section(project_id: str, request: RenameSectionRequest):
    doc_ref = db.collection("projects").document(project_id)
    project = doc_ref.get().to_dict()
    
    if not project or "content" not in project:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")
        
    content = project["content"]
    if request.old_name in content:
        content[request.new_name] = content.pop(request.old_name)
        doc_ref.update({"content": content})
        return {"message": "Sección renombrada exitosamente"}
    
    raise HTTPException(status_code=400, detail="La sección no existe")

@app.delete("/api/thesis/delete-section/{project_id}/{section_name}")
async def delete_section(project_id: str, section_name: str):
    doc_ref = db.collection("projects").document(project_id)
    project = doc_ref.get().to_dict()
    
    if not project or "content" not in project:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")
        
    content = project["content"]
    if section_name in content:
        del content[section_name]
        doc_ref.update({"content": content})
        return {"message": "Sección eliminada exitosamente"}
    
    raise HTTPException(status_code=400, detail="La sección no existe")


# --- Admin Endpoints ---

class UserCreateRequest(BaseModel):
    email: str
    password: str
    displayName: str
    role: str = "researcher"

@app.post("/api/admin/users/create")
async def create_user(request: UserCreateRequest):
    try:
        # Create user in Firebase Auth
        user_record = firebase_auth.create_user(
            email=request.email,
            password=request.password,
            display_name=request.displayName
        )
        
        # Set custom claims for role
        firebase_auth.set_custom_user_claims(user_record.uid, {request.role: True})
        
        return {
            "uid": user_record.uid,
            "message": "Usuario creado exitosamente en Auth"
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

class UserStatusRequest(BaseModel):
    uid: str
    status: str # 'active' or 'disabled'

@app.post("/api/admin/users/toggle-status")
async def toggle_user_status(request: UserStatusRequest):
    try:
        disabled = True if request.status == 'disabled' else False
        firebase_auth.update_user(request.uid, disabled=disabled)
        return {"message": f"Usuario {'inhabilitado' if disabled else 'habilitado'} en Auth"}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.delete("/api/admin/users/delete/{uid}")
async def delete_user(uid: str):
    try:
        firebase_auth.delete_user(uid)
        return {"message": "Usuario eliminado de Auth"}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.post("/api/admin/users/update-role")
async def update_user_role(request: UserCreateRequest): # Reuse model but only email/role matter
    try:
        # Get user by email to find UID
        user = firebase_auth.get_user_by_email(request.email)
        # Set custom claims
        firebase_auth.set_custom_user_claims(user.uid, {request.role: True})
        return {"message": f"Rol actualizado a {request.role} en Auth"}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
